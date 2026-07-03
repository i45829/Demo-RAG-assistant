"""
app.py — Демо-система «Литературный обзор по теме» для презентации руководству.

ЧТО ДЕЛАЕТ:
  1. Пользователь вводит тему в браузере (никакой установки — просто веб-страница).
  2. Приложение ищет источники через бесплатные научные API (OpenAlex, arXiv,
     Europe PMC, DOAJ, Crossref + Unpaywall для платных статей).
  3. Собранные данные передаются в LLM (Claude или Gemini — на выбор), которая
     синтезирует обзор литературы с нумерованными ссылками [n].
  4. Результат показывается в браузере и скачивается как .docx.

ЗАПУСК ЛОКАЛЬНО (для проверки перед деплоем, не обязателен для демо):
  pip install streamlit httpx python-docx
  streamlit run app.py

ДЕПЛОЙ (бесплатно, без установки на ПК начальства) — см. файл
deploy_streamlit_cloud.md в комплекте.
"""

import io
import os
import time
import base64
import json
from datetime import date

import httpx
import streamlit as st
from docx import Document
from docx.shared import Pt

# ============================== НАСТРОЙКИ СТРАНИЦЫ ==============================

st.set_page_config(page_title="Литературный обзор — демо", page_icon="📚", layout="wide")

# Значение по умолчанию; перезаписывается реальным email из сайдбара при запуске —
# так запросы к API попадают в "вежливый пул" (polite pool) и реже получают отказы.
CONTACT_EMAIL = "demo@example.com"


class _RetryableHTTPError(Exception):
    """Понятная ошибка с кодом статуса — чтобы предупреждения в интерфейсе были
    диагностируемыми, а не просто 'не удалось получить данные'."""
    def __init__(self, status_code, url):
        self.status_code = status_code
        super().__init__(f"HTTP {status_code} от {url}")


def _request_with_retry(method, url, retries=3, backoff=1.5, **kwargs):
    """Обёртка над httpx с повторными попытками для временных сбоев:
    429 (превышен лимит запросов — САМАЯ частая причина отказов у бесплатных API
    при вызовах с общего IP облачного хостинга, как у Streamlit Cloud), 502/503/504,
    таймауты. Для 429 уважает заголовок Retry-After, если он есть.
    follow_redirects=True по умолчанию — некоторые API (например arXiv) отдают 301
    с http на https, а httpx по умолчанию редиректы НЕ проходит и падает с ошибкой."""
    kwargs.setdefault("follow_redirects", True)
    headers = kwargs.get("headers") or {}
    headers.setdefault("User-Agent", f"lit-review-demo/1.0 (mailto:{CONTACT_EMAIL})")
    kwargs["headers"] = headers
    last_exc = None
    for attempt in range(retries + 1):
        try:
            r = getattr(httpx, method)(url, **kwargs)
            if r.status_code == 429 and attempt < retries:
                wait = float(r.headers.get("Retry-After", backoff * (attempt + 1) * 2))
                time.sleep(min(wait, 15))
                continue
            if r.status_code in (502, 503, 504) and attempt < retries:
                time.sleep(backoff * (attempt + 1))
                continue
            if r.status_code >= 400:
                last_exc = _RetryableHTTPError(r.status_code, url)
                raise last_exc
            return r
        except (httpx.TimeoutException, httpx.TransportError) as exc:
            last_exc = exc
            if attempt < retries:
                time.sleep(backoff * (attempt + 1))
                continue
            raise
        except _RetryableHTTPError:
            raise
    raise last_exc


def _get(url, **kwargs):
    return _request_with_retry("get", url, **kwargs)


def _post(url, **kwargs):
    return _request_with_retry("post", url, **kwargs)


# ============================== ПОИСК ПО ОТКРЫТЫМ API ==============================


def _trim(text, limit=1500):
    text = (text or "").strip().replace("\n", " ")
    return text if len(text) <= limit else text[:limit] + "…"




def search_pubmed(query, max_results=5):
    """PubMed E-utilities (NCBI) — 36M+ публикаций в медицине, биологии, химии, фармакологии.
    Без ключа (ключ увеличивает лимит до 10 запросов/сек, с ним до 100K/день).
    Отличие от Europe PMC: другая поисковая система, часто комплементарные результаты."""
    try:
        # Шаг 1: поиск PMID
        r = _get(
            "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi",
            params={"db": "pubmed", "term": query, "retmax": max_results,
                    "retmode": "json", "sort": "relevance"},
            timeout=20,
        )
        r.raise_for_status()
        id_list = r.json().get("esearchresult", {}).get("idlist", [])
        if not id_list:
            return []
        # Шаг 2: детали по PMID
        r2 = _get(
            "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi",
            params={"db": "pubmed", "id": ",".join(id_list), "retmode": "json"},
            timeout=20,
        )
        r2.raise_for_status()
        result_data = r2.json().get("result", {})
        # Шаг 3: аннотации
        r3 = _get(
            "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi",
            params={"db": "pubmed", "id": ",".join(id_list), "rettype": "abstract", "retmode": "xml"},
            timeout=25,
        )
        r3.raise_for_status()
        # Простое извлечение аннотаций из XML без lxml
        abstracts = {}
        import re as _re
        for m in _re.finditer(r'<PMID[^>]*>(\d+)</PMID>.*?<AbstractText[^>]*>(.*?)</AbstractText>',
                              r3.text, _re.DOTALL):
            pmid, abst = m.group(1), _re.sub(r'<[^>]+>', '', m.group(2))
            if pmid not in abstracts:
                abstracts[pmid] = abst.strip()
        items = []
        for pmid in id_list:
            meta = result_data.get(pmid, {})
            title = meta.get("title", "(без названия)")
            year = meta.get("pubdate", "н/д")[:4]
            authors = "; ".join(
                f"{a.get('name','')}" for a in (meta.get("authors") or [])[:4]
            ) or "н/д"
            items.append({
                "title": title,
                "authors": authors,
                "year": year,
                "source": "PubMed (NCBI)",
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
                "doi": "",
                "abstract": _trim(abstracts.get(pmid, "")),
            })
        return items
    except Exception as exc:
        st.warning(f"PubMed: не удалось получить данные ({exc})")
        return []


def search_arxiv(query, max_results=5):
    try:
        import feedparser
        r = _get(
            "https://export.arxiv.org/api/query",
            params={"search_query": f"all:{query}", "start": 0,
                    "max_results": max_results, "sortBy": "relevance"},
            timeout=20,
        )
        r.raise_for_status()
        feed = feedparser.parse(r.text)
        items = []
        for e in feed.entries:
            items.append({
                "title": _trim(getattr(e, "title", "(без названия)"), 300),
                "authors": ", ".join(a.name for a in getattr(e, "authors", [])[:5]) or "н/д",
                "year": (getattr(e, "published", "н/д") or "н/д")[:4],
                "source": "arXiv",
                "url": getattr(e, "link", ""),
                "doi": "",
                "abstract": _trim(getattr(e, "summary", "")),
            })
        return items
    except Exception as exc:
        st.warning(f"arXiv: не удалось получить данные ({exc})")
        return []


def search_europepmc(query, max_results=5):
    try:
        r = _get(
            "https://www.ebi.ac.uk/europepmc/webservices/rest/search",
            params={"query": query, "format": "json", "resultType": "core",
                    "pageSize": max_results, "email": CONTACT_EMAIL},
            timeout=20,
        )
        r.raise_for_status()
        items = []
        for res in (r.json().get("resultList") or {}).get("result", []):
            doi = res.get("doi") or ""
            items.append({
                "title": res.get("title") or "(без названия)",
                "authors": res.get("authorString") or "н/д",
                "year": res.get("pubYear") or "н/д",
                "source": "Europe PMC",
                "url": f"https://doi.org/{doi}" if doi else "",
                "doi": doi,
                "abstract": _trim(res.get("abstractText") or ""),
            })
        return items
    except Exception as exc:
        st.warning(f"Europe PMC: не удалось получить данные ({exc})")
        return []


def search_doaj(query, max_results=5):
    try:
        from urllib.parse import quote
        r = _get(
            f"https://doaj.org/api/search/articles/{quote(query)}",
            params={"page_size": max_results}, timeout=20,
        )
        r.raise_for_status()
        items = []
        for res in r.json().get("results", []):
            bib = res.get("bibjson", {})
            doi = next((i.get("id", "") for i in bib.get("identifier", [])
                        if (i.get("type") or "").lower() == "doi"), "")
            link = next((l.get("url", "") for l in bib.get("link", [])
                         if l.get("type") == "fulltext"), "") or (
                f"https://doi.org/{doi}" if doi else "")
            items.append({
                "title": bib.get("title") or "(без названия)",
                "authors": ", ".join(a.get("name", "") for a in (bib.get("author") or [])[:5]) or "н/д",
                "year": bib.get("year") or "н/д",
                "source": "DOAJ",
                "url": link,
                "doi": doi,
                "abstract": _trim(bib.get("abstract") or ""),
            })
        return items
    except Exception as exc:
        st.warning(f"DOAJ: не удалось получить данные ({exc})")
        return []


def search_semanticscholar(query, max_results=5):
    """Semantic Scholar — 200M+ работ. Без ключа лимит ~1 req/s с shared IP (Streamlit Cloud
    это бьёт). С бесплатным ключом лимит ~10 req/s. Ключ: semanticscholar.org/product/api.
    Ключ читается из переменной SS_API_KEY или из session_state['ss_api_key']."""
    ss_key = (os.environ.get("SS_API_KEY") or
              st.session_state.get("ss_api_key", "")).strip()
    headers = {}
    if ss_key:
        headers["x-api-key"] = ss_key
    try:
        r = _get(
            "https://api.semanticscholar.org/graph/v1/paper/search",
            params={
                "query": query, "limit": max_results,
                "fields": "title,year,authors,abstract,externalIds,venue,url",
            },
            headers=headers,
            timeout=20,
        )
        r.raise_for_status()
        items = []
        for p in (r.json().get("data") or []):
            doi = (p.get("externalIds") or {}).get("DOI", "")
            items.append({
                "title": p.get("title") or "(без названия)",
                "authors": ", ".join(a.get("name", "") for a in (p.get("authors") or [])[:5]) or "н/д",
                "year": p.get("year") or "н/д",
                "source": "Semantic Scholar",
                "url": f"https://doi.org/{doi}" if doi else (p.get("url") or ""),
                "doi": doi,
                "abstract": _trim(p.get("abstract") or ""),
            })
        return items
    except Exception as exc:
        st.warning(f"Semantic Scholar: не удалось получить данные ({exc})"
                   + ("" if ss_key else " — попробуйте добавить бесплатный ключ S2 в настройки источников"))
        return []


def search_openaire(query, max_results=5):
    """OpenAIRE Graph API — актуальный стабильный endpoint researchProducts.
    Старый /search/publications часто отдаёт 400 из-за смены схемы параметров."""
    try:
        r = _get(
            "https://api.openaire.eu/graph/v1/researchProducts",
            params={"search": query, "type": "publication", "pageSize": max_results},
            timeout=25,
        )
        r.raise_for_status()
        data = r.json()
        results = data.get("results", []) or []
        items = []
        for res in results[:max_results]:
            # v1 Graph API: плоская структура с mainTitle, publicationDate, authors, pids, description
            title = res.get("mainTitle") or "(без названия)"
            year = str(res.get("publicationDate") or "н/д")[:4]
            authors_raw = res.get("authors", []) or []
            authors = ", ".join(
                a.get("fullName", "") for a in authors_raw[:5] if isinstance(a, dict)
            ) or "н/д"
            pids = res.get("pids", []) or []
            doi = ""
            for pid in pids:
                if isinstance(pid, dict) and (pid.get("scheme") or "").lower() == "doi":
                    doi = pid.get("value", "")
                    break
            best_url = f"https://doi.org/{doi}" if doi else (res.get("id", "") or "")
            descriptions = res.get("descriptions", []) or []
            desc_text = ""
            if descriptions:
                first = descriptions[0]
                desc_text = first.get("value", "") if isinstance(first, dict) else str(first)
            abstract = _trim(desc_text)
            items.append({
                "title": title, "authors": authors, "year": year,
                "source": "OpenAIRE", "url": best_url, "doi": doi, "abstract": abstract,
            })
        return items
    except Exception as exc:
        st.warning(f"OpenAIRE: не удалось получить данные ({exc})")
        return []


def search_crossref(query, max_results=5):
    """Метаданные почти любого DOI, включая закрытые журналы (ScienceDirect и т.п.)."""
    try:
        r = _get(
            "https://api.crossref.org/works",
            params={"query": query, "rows": max_results, "mailto": CONTACT_EMAIL},
            timeout=20,
        )
        r.raise_for_status()
        items = []
        for it in (r.json().get("message") or {}).get("items", []):
            doi = it.get("DOI", "")
            year = None
            for k in ("published-print", "published-online", "published"):
                dp = (it.get(k) or {}).get("date-parts")
                if dp:
                    year = dp[0][0]
                    break
            items.append({
                "title": (it.get("title") or ["(без названия)"])[0],
                "authors": ", ".join(
                    f"{a.get('given','')} {a.get('family','')}".strip()
                    for a in (it.get("author") or [])[:5]
                ) or "н/д",
                "year": year or "н/д",
                "source": f"Crossref ({it.get('publisher', 'н/д')})",
                "url": f"https://doi.org/{doi}" if doi else "",
                "doi": doi,
                "abstract": "",  # Crossref обычно не даёт аннотацию
            })
        return items
    except Exception as exc:
        st.warning(f"Crossref: не удалось получить данные ({exc})")
        return []


def check_unpaywall(doi):
    """По DOI проверяет легальную бесплатную копию платной статьи."""
    if not doi:
        return None
    try:
        r = _get(f"https://api.unpaywall.org/v2/{doi}",
                       params={"email": CONTACT_EMAIL}, timeout=15)
        r.raise_for_status()
        data = r.json()
        if not data.get("is_oa"):
            return None
        best = data.get("best_oa_location") or {}
        return best.get("url_for_pdf") or best.get("url")
    except Exception:
        return None


# ============================== ПАТЕНТНЫЕ БАЗЫ (нужны свои ключи) ==============================
# Каждая функция возвращает (items, error) — error=None при успехе, иначе строка с
# понятным описанием проблемы (модель показывает её пользователю, а не падает).


# ============================== ПАТЕНТНЫЕ БАЗЫ (БЕЗ КЛЮЧЕЙ) ==============================



def search_core_oa(query, max_results=5):
    """CORE — крупнейший агрегатор открытых полных текстов (300M+ документов).
    Включает часть патентов EPO и отчёты. Бесплатный ключ CORE_API_KEY (core.ac.uk)."""
    api_key = os.environ.get("CORE_API_KEY", "")
    if not api_key:
        st.info("CORE: для поиска задайте CORE_API_KEY (бесплатно на core.ac.uk/services/api)")
        return []
    try:
        r = _get(
            "https://api.core.ac.uk/v3/search/outputs",
            params={"q": query, "limit": max_results},
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=25,
        )
        r.raise_for_status()
        items = []
        for d in (r.json().get("results") or [])[:max_results]:
            doi = d.get("doi") or ""
            items.append({
                "title": d.get("title") or "(без названия)",
                "authors": ", ".join(a.get("name","") for a in (d.get("authors") or [])[:4]) or "н/д",
                "year": str(d.get("yearPublished") or "н/д"),
                "source": "CORE",
                "url": d.get("downloadUrl") or (f"https://doi.org/{doi}" if doi else ""),
                "doi": doi,
                "abstract": _trim(d.get("abstract") or ""),
            })
        return items
    except Exception as exc:
        st.warning(f"CORE: не удалось получить данные ({exc})")
        return []





def search_rospatent(query, max_results=5, api_key=""):
    """Роспатент — Поисковая платформа (searchplatform.rospatent.gov.ru).
    Требует JWT-токен (API-ключ), выдаётся в личном кабинете Роспатента.
    Поисковая база: патенты РФ, СНГ, PCT-минимум (CN, US, EP, JP, KR, DE, FR и др.)."""
    if not api_key or not api_key.strip():
        return [], "Роспатент: не указан API-ключ (JWT) — источник пропущен."
    try:
        r = _post(
            "https://searchplatform.rospatent.gov.ru/patsearch/v0.2/search",
            headers={
                "Authorization": f"Bearer {api_key.strip()}",
                "Content-Type": "application/json",
            },
            json={"qn": query, "limit": max_results},
            timeout=30,
        )
        r.raise_for_status()
        data = r.json()
    except Exception as exc:
        return [], f"Роспатент: ошибка запроса ({exc})"

    items = []
    for hit in (data.get("hits") or [])[:max_results]:
        biblio_ru = (hit.get("biblio") or {}).get("ru", {}) or {}
        biblio_en = (hit.get("biblio") or {}).get("en", {}) or {}
        title = biblio_ru.get("title") or biblio_en.get("title") or "(без названия)"
        inventors = biblio_ru.get("inventor") or biblio_en.get("inventor") or []
        authors = ", ".join(inv.get("name", "") for inv in inventors[:5]) or "н/д"
        common = hit.get("common") or {}
        pub_date = common.get("publication_date", "н/д")
        year = pub_date[:4] if pub_date and pub_date != "н/д" else "н/д"
        country = common.get("publishing_office", "")
        doc_num = common.get("document_number", "")
        kind = common.get("kind", "")
        pn = f"{country}{doc_num}{kind}" if country and doc_num else hit.get("id", "")
        snippet = hit.get("snippet") or {}
        descr = snippet.get("description", "")
        # Очищаем HTML-тэги подсветки
        import re as _re
        descr_clean = _re.sub(r"<[^>]+>", "", descr) if descr else ""
        items.append({
            "title": title,
            "authors": authors,
            "year": year,
            "source": "Роспатент",
            "url": f"https://searchplatform.rospatent.gov.ru/docs/{hit.get('id', '')}" if hit.get("id") else "",
            "doi": "",
            "abstract": _trim(descr_clean),
        })
    return items, None


PATENT_SOURCES = {
    "Роспатент (патенты РФ, СНГ, PCT — нужен JWT-ключ)": (
        search_rospatent,
        [("API-ключ Роспатента (JWT-токен)", "rospatent_key")],
    ),
}



# PatentsView возвращает (items, None) / ([], msg) — адаптируем в обёртку
def _patentsview_wrap(query, max_results=5):
    """Обёртка, чтобы PatentsView вписывался в обычный SOURCES (возвращает list)."""
    result = search_patentsview_legacy(query, max_results)
    return result  # search_patentsview_legacy уже возвращает list и сам вызывает st.warning

SOURCES = {
    "PubMed / NCBI (медицина, химия, биология)": search_pubmed,
    "Europe PMC (PMC + Европейская коллекция)": search_europepmc,
    "OpenAIRE (85M+ открытых публикаций, без ключа)": search_openaire,
    "Crossref (метаданные DOI + Unpaywall)": search_crossref,
    "DOAJ (открытые журналы)": search_doaj,
    "arXiv (препринты, тема на англ.)": search_arxiv,
}

# ============================== ИЗВЛЕЧЕНИЕ ПОИСКОВЫХ ЗАПРОСОВ ==============================
# Источники, для которых нужен английский запрос (они плохо ищут по русским словам)
ENGLISH_CENTRIC = {
    "PubMed / NCBI (медицина, химия, биология)",
    "OpenAIRE (85M+ открытых публикаций, без ключа)",
    "arXiv (препринты, тема на англ.)",
    "DOAJ (открытые журналы)",
}


def translate_topic(topic: str, api_key: str, provider: str, model_id: str) -> dict:
    """Переводит тему на русский и английский (сохраняет исходный смысл целиком).
    В отличие от прошлой версии — не пытается «извлечь ключевые термины», потому что
    это выкидывает контекст (при теме «Сверхкритическая вода. Тренды 2023-2026»
    извлечение оставляло только «сверхкритическая вода», теряя аспект «тренды»)."""
    system = (
        "Ты — переводчик. Переведи тему научного исследования на русский и английский. "
        "Сохрани ВСЕ смысловые аспекты (тренды, обзор, конкретные вещества и т.п.), "
        "убери только годы/даты, если они есть. "
        "Верни ТОЛЬКО JSON без Markdown, без пояснений. "
        'Схема: {"ru": "тема на русском", "en": "topic in English"}. '
        "Примеры: "
        'тема «Сверхкритическая вода. Тренды 2023-2026» -> '
        '{"ru": "Сверхкритическая вода, современные тренды исследований", '
        '"en": "Supercritical water, current research trends"}; '
        'тема «Масляные дисперсии как препаративная форма для СЗР» -> '
        '{"ru": "Масляные дисперсии как препаративная форма для средств защиты растений", '
        '"en": "Oil dispersions as formulation type for plant protection products"}. '
        "ТОЛЬКО JSON."
    )
    try:
        raw = call_llm(provider, model_id, system, topic, api_key, temperature=0.0)
        import json as _json, re as _re
        clean = _re.sub(r"```[a-z]*", "", raw).strip().strip("`")
        parsed = _json.loads(clean)
        ru = str(parsed.get("ru", topic)).strip() or topic
        en = str(parsed.get("en", topic)).strip() or topic
        return {"ru": ru, "en": en}
    except Exception:
        # Фолбэк: используем оригинальную тему в обоих полях (лучше, чем ничего)
        return {"ru": topic, "en": topic}


# ============================== ФИЛЬТР РЕЛЕВАНТНОСТИ (embeddings) ==============================

def _gemini_embed(texts: list, api_key: str, model_id: str,
                  task_type: str = "SEMANTIC_SIMILARITY") -> list:
    """Пакетный вызов Gemini embedContent для списка текстов.
    model_id — например 'text-embedding-004' или 'gemini-embedding-001'."""
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_id}:batchEmbedContents"
    body = {
        "requests": [
            {
                "model": f"models/{model_id}",
                "content": {"parts": [{"text": t}]},
                "taskType": task_type,
            }
            for t in texts
        ]
    }
    r = httpx.post(url, params={"key": api_key}, json=body, timeout=30)
    r.raise_for_status()
    data = r.json()
    return [emb["values"] for emb in data.get("embeddings", [])]


def _polza_embed(texts: list, api_key: str, model_id: str) -> list:
    """Polza.ai embedding endpoint — OpenAI-совместимый.
    model_id — например 'openai/text-embedding-3-large' или 'google/gemini-embedding-001'."""
    r = httpx.post(
        "https://polza.ai/api/v1/embeddings",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={"model": model_id, "input": texts},
        timeout=30,
    )
    r.raise_for_status()
    data = r.json()
    # Ответ соответствует OpenAI: data.data — массив с полями embedding и index
    result = sorted(data.get("data", []), key=lambda x: x.get("index", 0))
    return [item["embedding"] for item in result]


def _embed(provider: str, texts: list, api_key: str, model_id: str) -> list:
    """Универсальный dispatcher по провайдеру эмбеддингов."""
    if provider == "Google Gemini":
        return _gemini_embed(texts, api_key, model_id)
    return _polza_embed(texts, api_key, model_id)


def _cosine_similarity(v1: list, v2: list) -> float:
    """Косинусная близость двух векторов (без numpy — чистый Python)."""
    dot = sum(a * b for a, b in zip(v1, v2))
    n1 = sum(a * a for a in v1) ** 0.5
    n2 = sum(b * b for b in v2) ** 0.5
    if n1 == 0 or n2 == 0:
        return 0.0
    return dot / (n1 * n2)


def filter_by_relevance(topic_translated: dict, items: list, provider: str,
                        api_key: str, embedding_model: str,
                        threshold: float = 0.55) -> tuple:
    """Оставляет только статьи, косинусно близкие к теме (по эмбеддингам).
    Возвращает (отфильтрованный список, оценки для отчёта).

    threshold=0.55 подобран как разумный компромисс:
    - <0.45: мусор проходит
    - 0.55: отсекает явно нерелевантное, сохраняя пограничные случаи
    - >0.65: слишком строго, теряются близкие по духу работы"""
    if not items:
        return items, []

    # Собираем «текст статьи» — заголовок + аннотация, если есть
    def _article_text(it):
        parts = [it.get("title", "")]
        if it.get("abstract"):
            parts.append(it["abstract"])
        return " ".join(parts)[:2000]  # ограничение по токенам эмбеддингов

    # Тексты для эмбеддинга: 2 темы (ru/en) + N статей
    topic_texts = [topic_translated["ru"], topic_translated["en"]]
    article_texts = [_article_text(it) for it in items]
    all_texts = topic_texts + article_texts

    try:
        vectors = _embed(provider, all_texts, api_key, embedding_model)
    except Exception as exc:
        st.warning(
            f"⚠️ Не удалось вычислить эмбеддинги ({exc}). "
            f"Фильтрация по релевантности пропущена, все статьи оставлены."
        )
        return items, []

    if len(vectors) < 2 + len(items):
        return items, []

    topic_vec_ru = vectors[0]
    topic_vec_en = vectors[1]
    article_vecs = vectors[2:]

    # Для каждой статьи берём максимум из близости к RU и EN версии темы
    scored = []
    for it, avec in zip(items, article_vecs):
        sim_ru = _cosine_similarity(topic_vec_ru, avec)
        sim_en = _cosine_similarity(topic_vec_en, avec)
        score = max(sim_ru, sim_en)
        scored.append((it, score))

    # Сортируем по убыванию релевантности
    scored.sort(key=lambda x: x[1], reverse=True)

    # Оставляем прошедшие порог, но не меньше 3 (даже если все близки к границе)
    filtered = [it for it, score in scored if score >= threshold]
    if len(filtered) < 3 and len(scored) >= 3:
        filtered = [it for it, _ in scored[:3]]

    return filtered, scored


# ============================== СИНТЕЗ ОБЗОРА (LLM) ==============================

REVIEW_SYSTEM_PROMPT = """Ты — научный аналитик, готовящий раздел «Обзор литературы» для диссертации по естественнонаучной/технической теме.

ВАЖНО: Источники уже прошли автоматическую фильтрацию по релевантности (эмбеддинги + косинусная близость), поэтому все переданные тебе статьи заведомо связаны с темой. Твоя задача — не фильтровать, а СИНТЕЗИРОВАТЬ содержание.

ШАГ 1 — ДОПОЛНИТЕЛЬНАЯ ПРОВЕРКА (только для явных сбоев):
Если конкретная статья очевидно не про тему (например, экономика/социология при теме по физической химии) — молча исключи её. В обычных случаях всё оставляй.

ШАГ 2 — НАПИСАНИЕ ОБЗОРА:
Пиши только то, что реально есть в аннотациях релевантных источников. Извлекай:
- точные названия веществ, материалов, соединений, реагентов, методов, приборов;
- химические формулы и обозначения — буквально как в источнике;
- числа: температуры, давления, концентрации, pH, эффективность, проценты, p-значения;
- механизмы и принципы действия — как описаны в источнике;
- конкретные результаты экспериментов с числами.

ЗАПРЕЩЕНО в теле обзора:
- «было найдено X источников», «источники были проанализированы», «проведён поиск»;
- обобщения без цифр и ссылок; пересказ того, что ты делал.

ГАРАНТИЯ НЕПУСТОГО ВЫВОДА:
- Если после фильтрации осталось ≥1 релевантного источника — напиши полноценный обзор по нему/ним.
- Если релевантных источников не осталось совсем — напиши в разделе «Обзор литературы» одно честное предложение об этом, затем в разделе «Выводы» перечисли, что именно нужно искать и в каких базах.
- Никогда не возвращай пустой текст.

Каждое фактическое утверждение — со ссылкой [n] на номер источника.

Структура ответа (строго Markdown, на русском языке):
## Введение
2–4 предложения: тема, актуальность, границы обзора. Без канцелярских клише.

## Обзор литературы
Тематический синтез конкретных данных из источников. Ссылка [n] на каждый факт. Подзаголовки ### по подтемам при необходимости.

## Выводы и направления дальнейших исследований
- Обоснованные выводы из данных [n]
- Пробелы в изученности темы (каких данных не хватает)
- Рекомендации (явно помечены как предположения)

## Список литературы
[n] Авторы. Название. Год. Источник. URL."""


# Максимум символов на один источник зависит от числа источников:
# мало источников → длинные аннотации; много → обрезаем, чтобы не перегрузить контекст.
_MAX_TOTAL_CONTEXT = 60_000   # ~15K токенов — комфортно для всех моделей
_MIN_ABSTRACT_PER_ITEM = 300  # минимум даже при большом количестве источников


def build_user_message(topic, items):
    n_with_abstract = sum(1 for it in items if it.get("abstract", "").strip())
    # Динамический лимит аннотации на источник
    if items:
        budget = max(_MIN_ABSTRACT_PER_ITEM,
                     (_MAX_TOTAL_CONTEXT - len(topic) - 500) // len(items))
    else:
        budget = 1500
    lines = [
        f"ТЕМА ОБЗОРА: {topic}", "",
        f"ИСТОЧНИКОВ ПЕРЕДАНО: {len(items)}, из них с аннотацией: {n_with_abstract}.",
        "Если аннотация пуста ('н/д') — используй название/авторов/год,",
        "пометив '(аннотация недоступна)'. НЕ оставляй обзор пустым.", "",
    ]
    for i, it in enumerate(items, 1):
        abst = (it.get("abstract") or "н/д").strip()
        if len(abst) > budget:
            abst = abst[:budget] + "…"
        lines.append(
            f"[{i}] {it['title']} ({it['year']}). Авторы: {it['authors']}. "
            f"Источник: {it['source']}. URL: {it['url'] or 'н/д'}\n"
            f"Аннотация: {abst}\n"
        )
    return "\n".join(lines)


def call_gemini(system_prompt, user_message, api_key, model="gemini-2.0-flash", temperature=0.3):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    body = {
        "system_instruction": {"parts": [{"text": system_prompt}]},
        "contents": [{"role": "user", "parts": [{"text": user_message}]}],
        "generationConfig": {"temperature": temperature},
    }
    last_exc = None
    for attempt in range(3):
        try:
            r = httpx.post(url, params={"key": api_key}, json=body, timeout=120)
            if r.status_code == 429:
                wait = float(r.headers.get("Retry-After", 10 * (attempt + 1)))
                time.sleep(min(wait, 30))
                continue
            if r.status_code in (502, 503, 504) and attempt < 2:
                time.sleep(5 * (attempt + 1))
                continue
            r.raise_for_status()
            data = r.json()
            candidate = data.get("candidates", [{}])[0]
            finish = candidate.get("finishReason", "STOP")
            if finish == "SAFETY":
                raise ValueError(
                    "Gemini заблокировал ответ фильтром безопасности (finishReason=SAFETY). "
                    "Попробуйте переформулировать тему или выбрать другую модель."
                )
            if finish == "MAX_TOKENS":
                # Частичный ответ — вернём что есть, дальше в коде проверим длину
                parts = (candidate.get("content") or {}).get("parts", [])
                return parts[0].get("text", "") if parts else ""
            parts = (candidate.get("content") or {}).get("parts", [])
            text = parts[0].get("text", "") if parts else ""
            return text
        except (httpx.TimeoutException, httpx.TransportError) as exc:
            last_exc = exc
            if attempt < 2:
                time.sleep(5 * (attempt + 1))
                continue
            raise
        except Exception as exc:
            last_exc = exc
            if attempt < 2 and "503" in str(exc):
                time.sleep(5 * (attempt + 1))
                continue
            raise
    raise last_exc


def call_polza(system_prompt, user_message, api_key, model, temperature=0.3):
    """Polza.ai — агрегатор моделей с OpenAI-совместимым API."""
    body = {
        "model": model, "temperature": temperature,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
    }
    last_exc = None
    for attempt in range(3):
        try:
            r = httpx.post(
                "https://polza.ai/api/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json=body, timeout=120,
            )
            if r.status_code == 429:
                wait = float(r.headers.get("Retry-After", 10 * (attempt + 1)))
                time.sleep(min(wait, 30))
                continue
            if r.status_code in (502, 503, 504) and attempt < 2:
                time.sleep(5 * (attempt + 1))
                continue
            r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"]
        except (httpx.TimeoutException, httpx.TransportError) as exc:
            last_exc = exc
            if attempt < 2:
                time.sleep(5 * (attempt + 1))
                continue
            raise
        except Exception as exc:
            last_exc = exc
            if attempt < 2 and "503" in str(exc):
                time.sleep(5 * (attempt + 1))
                continue
            raise
    raise last_exc


def call_llm(provider, model_id, system_prompt, user_message, api_key, temperature=0.3):
    if provider == "Google Gemini":
        return call_gemini(system_prompt, user_message, api_key, model=model_id, temperature=temperature)
    return call_polza(system_prompt, user_message, api_key, model=model_id, temperature=temperature)


# ------------------------- Загрузка списка доступных моделей -------------------------


@st.cache_data(ttl=600, show_spinner=False)
def fetch_gemini_models(api_key):
    """Список моделей Gemini, поддерживающих generateContent."""
    r = httpx.get(
        "https://generativelanguage.googleapis.com/v1beta/models",
        params={"key": api_key}, timeout=20,
    )
    r.raise_for_status()
    out = []
    for m in r.json().get("models", []):
        if "generateContent" in (m.get("supportedGenerationMethods") or []):
            model_id = m["name"].removeprefix("models/")
            out.append((model_id, m.get("displayName") or model_id))
    return out


@st.cache_data(ttl=600, show_spinner=False)
def fetch_polza_models(api_key):
    """Список чат-моделей, доступных через Polza.ai (агрегатор OpenAI/Anthropic/Google и др.)."""
    r = httpx.get(
        "https://polza.ai/api/v1/models",
        params={"type": "chat"},
        headers={"Authorization": f"Bearer {api_key}"}, timeout=20,
    )
    r.raise_for_status()
    out = []
    for m in r.json().get("data", []):
        out.append((m["id"], m.get("name") or m["id"]))
    return out


@st.cache_data(ttl=600, show_spinner=False)
def fetch_polza_embedding_models(api_key):
    """Список embedding-моделей через Polza.ai (google/gemini-embedding-001,
    openai/text-embedding-3-large, qwen/qwen3-embedding-8b и др.)."""
    r = httpx.get(
        "https://polza.ai/api/v1/models",
        params={"type": "embedding"},
        headers={"Authorization": f"Bearer {api_key}"}, timeout=20,
    )
    r.raise_for_status()
    out = []
    for m in r.json().get("data", []):
        out.append((m["id"], m.get("name") or m["id"]))
    return out


@st.cache_data(ttl=600, show_spinner=False)
def fetch_gemini_embedding_models(api_key):
    """Список embedding-моделей у Google Gemini.
    Отфильтровываем модели, поддерживающие embedContent."""
    r = httpx.get(
        "https://generativelanguage.googleapis.com/v1beta/models",
        params={"key": api_key}, timeout=20,
    )
    r.raise_for_status()
    out = []
    for m in r.json().get("models", []):
        methods = m.get("supportedGenerationMethods") or []
        if "embedContent" in methods or "batchEmbedContents" in methods:
            model_id = m["name"].removeprefix("models/")
            out.append((model_id, m.get("displayName") or model_id))
    return out


def fetch_embedding_models(provider, api_key):
    if provider == "Google Gemini":
        return fetch_gemini_embedding_models(api_key)
    return fetch_polza_embedding_models(api_key)


def fetch_models(provider, api_key):
    if provider == "Google Gemini":
        return fetch_gemini_models(api_key)
    return fetch_polza_models(api_key)


# ============================== ЭКСПОРТ В WORD ==============================


def build_docx(topic, review_markdown):
    doc = Document()
    doc.add_heading("Литературный обзор", level=0)
    doc.add_paragraph(f"Тема: {topic}")
    doc.add_paragraph(f"Дата формирования: {date.today().strftime('%d.%m.%Y')}")
    doc.add_paragraph("Сгенерировано автоматически — демонстрационная версия.").italic = True

    for raw_line in review_markdown.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)
        else:
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ==================================== UI ====================================

st.title("📚 Демо: автоматический литературный обзор")
st.caption(
    "Вводите тему → система ищет источники в открытых научных базах → "
    "ИИ формирует обзор со ссылками → скачиваете Word-документ."
)

with st.sidebar:
    st.header("Настройки")

    provider = st.selectbox(
        "Провайдер LLM",
        ["Google Gemini", "Polza.ai (агрегатор моделей)"],
        help="Google Gemini — бесплатный тариф от Google. "
             "Polza.ai — платный (в рублях) агрегатор, даёт доступ к сотням моделей "
             "(OpenAI, Anthropic, Google и др.) через один ключ.",
    )
    key_label = "Google AI Studio" if provider == "Google Gemini" else "Polza.ai"
    api_key = st.text_input(
        f"API-ключ {key_label}",
                help="Ключ не сохраняется — используется только для этого запроса в этой сессии.",
    )

    model_id = None
    if api_key:
        try:
            with st.spinner("Загружаю список доступных моделей…"):
                model_options = fetch_models(provider, api_key)
        except Exception as exc:
            model_options = []
            st.error(f"Не удалось загрузить список моделей: {exc}")

        if model_options:
            labels = [f"{name}  ·  {mid}" if name != mid else mid for mid, name in model_options]
            idx = st.selectbox(
                "Модель", range(len(model_options)), format_func=lambda i: labels[i]
            )
            model_id = model_options[idx][0]
            st.caption(f"Доступно моделей: {len(model_options)}")
        else:
            st.warning("Список моделей пуст — проверьте ключ.")

        # Модель эмбеддингов для фильтра релевантности
        embedding_model_id = None
        try:
            with st.spinner("Загружаю список embedding-моделей…"):
                emb_options = fetch_embedding_models(provider, api_key)
        except Exception as exc:
            emb_options = []
            st.caption(f"⚠️ Embedding-модели недоступны: {exc}")

        if emb_options:
            emb_labels = [f"{name}  ·  {mid}" if name != mid else mid for mid, name in emb_options]
            # Пробуем найти дефолт — предпочитаем text-embedding-3-small или text-embedding-004
            default_idx = 0
            for pref in ("text-embedding-3-small", "text-embedding-004",
                         "openai/text-embedding-3-small", "google/gemini-embedding-001",
                         "gemini-embedding-001"):
                for i, (mid, _) in enumerate(emb_options):
                    if mid.endswith(pref) or mid == pref:
                        default_idx = i
                        break
                else:
                    continue
                break
            emb_idx = st.selectbox(
                "🎯 Модель эмбеддингов (фильтр релевантности)",
                range(len(emb_options)),
                index=default_idx,
                format_func=lambda i: emb_labels[i],
                help="Используется для оценки близости найденных статей к теме. "
                     "Меньшие модели быстрее и дешевле, большие — точнее."
            )
            embedding_model_id = emb_options[emb_idx][0]
    else:
        st.caption("Введите API-ключ, чтобы загрузить список моделей.")
        embedding_model_id = None

    st.divider()
    st.subheader("Источники для поиска")
    contact_email_input = st.text_input(
        "Email для научных API (необязательно)",
        placeholder="you@example.com",
        help="Некоторые API (OpenAlex, Crossref, Europe PMC) реже отказывают в "
             "обслуживании, если передавать реальный email — это их 'вежливый пул'. "
             "Ключ/пароль здесь не нужен, только адрес."
    )
    if contact_email_input.strip():
        CONTACT_EMAIL = contact_email_input.strip()

    all_source_names = list(SOURCES.keys()) + list(PATENT_SOURCES.keys())
    selected_sources = st.multiselect(
        "Выберите базы", all_source_names,
        default=list(SOURCES.keys()),
    )
    per_source = st.slider(
        "Релевантных источников на каждую базу", 2, 10, 4,
        help="Целевое число статей от каждой базы ПОСЛЕ фильтрации по релевантности. "
             "Приложение тянет из API в 3× больше и оставляет топ-N по эмбеддингам."
    )
    check_oa = st.checkbox(
        "Проверять открытый доступ (Unpaywall) для найденных DOI", value=True,
        help="Ищет легальную бесплатную копию, если статья из платного журнала."
    )

    # Поля для ключей выбранных патентных источников
    selected_patent_sources = [s for s in selected_sources if s in PATENT_SOURCES]
    patent_credentials = {}
    if selected_patent_sources:
        st.divider()
        st.subheader("Ключи патентных API")
        st.caption(
            "Для выбранных патентных баз нужен свой ключ. "
            "Без ключа источник будет пропущен при формировании обзора."
        )
        for name in selected_patent_sources:
            _, cred_fields = PATENT_SOURCES[name]
            with st.expander(name, expanded=True):
                values = {}
                for field_label, session_key in cred_fields:
                    values[session_key] = st.text_input(
                        field_label, key=f"cred_{session_key}",
                    )
                patent_credentials[name] = values


tab1, tab2 = st.tabs(["📚 Обзор", "⚙️ Настройка промта"])

with tab1:
    temperature = st.slider(
        "Температура ответа", min_value=0.0, max_value=1.0, value=0.5, step=0.05,
        help="Ниже (0.0–0.3) — более точный, предсказуемый, "
             "\"academic\" ответ. Выше (0.7–1.0) — больше вариативности и "
             "\"творчества\", но выше риск неточностей."
    )

    topic = st.text_area(
        "Тема литературного обзора",
        placeholder="Например: Барьеры перехода на электромобили в Европе (2023–2025): "
                    "инфраструктура, цена, меры господдержки",
        height=100,
    ).strip()

    go = st.button(
        "🔎 Сформировать обзор", type="primary",
        disabled=not (topic and api_key and model_id and selected_sources),
    )

    if not api_key:
        st.info("Введите API-ключ в боковой панели, чтобы начать.")
    elif not model_id:
        st.info("Выберите модель в боковой панели, чтобы начать.")

    if go:
        # Шаг 0: переводим тему на русский и английский (сохраняя весь смысл)
        with st.spinner("🌐 Перевожу тему на нужные языки для поиска…"):
            queries = translate_topic(topic, api_key, provider, model_id)

        with st.expander("🌐 Тема на языках поисковых баз", expanded=False):
            st.markdown(
                f"**Русский:** `{queries['ru']}`  \n"
                f"**English:** `{queries['en']}`  \n"
                "_Тема переведена целиком, все смысловые аспекты сохранены. "
                "Англоязычным базам уходит английский вариант, остальным — русский._"
            )

        all_items = []
        # Список кортежей (item, score, source_name) для итогового отчёта
        all_scored_report = []
        # Множитель добора: тянем в 3 раза больше, потом фильтруем
        OVERFETCH_MULTIPLIER = 3
        MAX_FETCH_PER_SOURCE = 30  # верхний предел — не бомбить API

        progress = st.progress(0.0, text="Поиск источников…")
        for i, name in enumerate(selected_sources):
            progress.progress((i) / len(selected_sources), text=f"Ищу в {name}…")
            api_query = queries["en"] if name in ENGLISH_CENTRIC else queries["ru"]

            # Тянем с запасом: per_source * 3, но не больше MAX_FETCH_PER_SOURCE
            fetch_count = min(per_source * OVERFETCH_MULTIPLIER, MAX_FETCH_PER_SOURCE)

            if name in SOURCES:
                found = SOURCES[name](api_query, fetch_count)
            elif name in PATENT_SOURCES:
                fn, cred_fields = PATENT_SOURCES[name]
                creds = patent_credentials.get(name, {})
                args = [creds.get(session_key, "") for _, session_key in cred_fields]
                found, err = fn(api_query, fetch_count, *args)
                if err:
                    st.warning(err)
            else:
                found = []

            # Фильтруем эту базу отдельно, добирая до per_source
            if found and embedding_model_id and len(found) > per_source:
                try:
                    _, scored_here = filter_by_relevance(
                        queries, found, provider, api_key, embedding_model_id,
                        threshold=0.0,  # порог отключён — берём топ-N по score
                    )
                    # scored_here уже отсортирован по убыванию score
                    top_items = [it for it, _ in scored_here[:per_source]]
                    all_items.extend(top_items)
                    # Сохраняем для отчёта (только оставленные)
                    for it, score in scored_here[:per_source]:
                        all_scored_report.append((it, score, name, True))
                    for it, score in scored_here[per_source:]:
                        all_scored_report.append((it, score, name, False))
                except Exception as exc:
                    st.warning(f"{name}: фильтр эмбеддингов не сработал ({exc}), "
                               f"беру первые {per_source} по порядку API")
                    all_items.extend(found[:per_source])
            else:
                # Либо мало источников, либо нет embedding-модели — берём как есть
                all_items.extend(found[:per_source])

        progress.progress(1.0, text="Поиск завершён")

        # Отчёт: показываем что взяли и что отбросили
        if all_scored_report:
            kept_count = sum(1 for *_, kept in all_scored_report if kept)
            dropped_count = len(all_scored_report) - kept_count
            with st.expander(
                f"🎯 Фильтр релевантности: оставлено {kept_count}, "
                f"отброшено {dropped_count} (по эмбеддингам)", expanded=False
            ):
                # Группируем по базе
                by_source = {}
                for it, score, source_name, kept in all_scored_report:
                    by_source.setdefault(source_name, []).append((it, score, kept))
                for source_name, entries in by_source.items():
                    st.markdown(f"**{source_name}**")
                    for it, score, kept in entries:
                        mark = "✅" if kept else "❌"
                        st.markdown(f"- {mark} `{score:.3f}` · {it['title'][:100]}")

        if check_oa:
            with st.spinner("Проверяю открытый доступ по DOI (Unpaywall)…"):
                for it in all_items:
                    if it.get("doi"):
                        oa_url = check_unpaywall(it["doi"])
                        if oa_url:
                            it["url"] = oa_url
                            it["source"] += " [открытая копия найдена]"

        if not all_items:
            st.error("Ничего не найдено. Попробуйте переформулировать тему или выбрать другие источники.")
            st.stop()

        with st.expander(f"📎 Найдено релевантных источников: {len(all_items)} (нажмите, чтобы посмотреть)"):
            for i, it in enumerate(all_items, 1):
                st.markdown(f"**[{i}] {it['title']}** ({it['year']}) — *{it['source']}*  \n{it['url']}")

        with st.spinner("ИИ формирует обзор литературы… (может занять до 1 мин)"):
            user_msg = build_user_message(topic, all_items)
            try:
                active_prompt = st.session_state.get("system_prompt", REVIEW_SYSTEM_PROMPT)
                review = call_llm(provider, model_id, active_prompt, user_msg, api_key, temperature=temperature)
            except Exception as exc:
                err_str = str(exc)
                if "503" in err_str:
                    st.error(
                        f"❌ Сервер модели временно недоступен (503) — попробуйте нажать "
                        f"«Сформировать обзор» ещё раз через 10–20 секунд. "
                        f"Это временный сбой на стороне провайдера, не проблема приложения."
                    )
                elif "429" in err_str:
                    st.error(
                        f"❌ Превышен лимит запросов к модели (429). "
                        f"Подождите 30–60 секунд и повторите."
                    )
                else:
                    st.error(f"❌ Ошибка обращения к модели: {exc}")
                st.stop()

        # Проверяем что обзор реально сформирован (не пустая строка)
        if not review or len(review.strip()) < 300:
            n_with_abstract = sum(1 for it in all_items if it.get("abstract","").strip())
            if n_with_abstract == 0:
                st.warning(
                    "⚠️ Источники найдены, но аннотации пустые. "
                    "Пробую сформировать обзор только по заголовкам..."
                )
            else:
                st.warning("⚠️ Модель вернула слишком короткий ответ. Пробую ещё раз...")
            # Авто-повтор: запасной промт + temperature 0.7
            FALLBACK_PROMPT = (
                "You are a scientific summarizer. Write a concise literature review in Russian "
                "based strictly on the provided sources. Include specific substances, formulas, "
                "quantities, and methods mentioned in abstracts. Use [n] citations. Structure: "
                "## Введение / ## Обзор литературы / ## Выводы / ## Список литературы. "
                "Even if abstracts are empty, write based on titles and authors. "
                "NEVER return an empty or very short response. Minimum 400 words."
            )
            try:
                with st.spinner("Повторная генерация с запасным промтом..."):
                    review = call_llm(
                        provider, model_id, FALLBACK_PROMPT,
                        build_user_message(topic, all_items),
                        api_key, temperature=0.7,
                    )
            except Exception as exc:
                review = ""
            if not review or len(review.strip()) < 200:
                st.error(
                    "❌ Модель не смогла сформировать обзор.\n\n"
                    f"Найдено источников: {len(all_items)}, с аннотацией: {n_with_abstract}.\n\n"
                    "**Возможные причины:**\n"
                    "- Тема слишком специфична и источники не найдены — попробуйте тему на английском\n"
                    "- Все найденные источники отфильтрованы как нерелевантные — выберите другие базы\n"
                    "- Gemini заблокировал запрос фильтром безопасности — попробуйте Polza.ai или другую модель"
                )
                st.stop()

        st.success("Обзор готов!")
        st.markdown(review)

        docx_buf = build_docx(topic, review)
        st.download_button(
            "⬇️ Скачать как Word (.docx)",
            data=docx_buf,
            file_name=f"literature_review_{date.today().isoformat()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    st.divider()
    st.caption(
        "Демо-версия. Источники: PubMed / NCBI, Europe PMC, OpenAIRE, Crossref, DOAJ, arXiv, Роспатент. "
        "Полный список используемых баз и логика проверки цитирования — в промт-шаблоне проекта."
    )

with tab2:
    st.subheader("Системный промт для генерации обзора")
    st.caption(
        "Здесь задаётся инструкция, по которой модель формирует обзор литературы. "
        "Можно скорректировать структуру, стиль, требования к цитированию и т.д."
    )

    if "system_prompt" not in st.session_state:
        st.session_state["system_prompt"] = REVIEW_SYSTEM_PROMPT

    edited_prompt = st.text_area(
        "Промт", value=st.session_state["system_prompt"], height=500,
        key="prompt_editor",
    )

    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Сохранить промт"):
            st.session_state["system_prompt"] = edited_prompt
            st.success("Промт сохранён и будет использован при следующей генерации.")
    with col2:
        if st.button("↩️ Сбросить к варианту по умолчанию"):
            st.session_state["system_prompt"] = REVIEW_SYSTEM_PROMPT
            st.rerun()
